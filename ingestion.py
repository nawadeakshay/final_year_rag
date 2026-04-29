"""
ingestion.py — Document loading, parsing, chunking, embedding, and FAISS indexing.

Pipeline:
  1. Discover all supported files in DOCUMENTS_DIR
  2. Parse with `unstructured` (preserves metadata: page, section, source)
  3. Chunk with RecursiveCharacterTextSplitter
  4. Embed with SentenceTransformers (normalised for cosine sim)
  5. Store in FAISS (persisted to disk)
  6. Build & persist BM25 index for sparse retrieval

Run directly:
  python ingestion.py
"""

import os
import pickle
import hashlib
from pathlib import Path
from typing import List, Tuple

import numpy as np
from loguru import logger
from tqdm import tqdm

# LangChain
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Embeddings
from sentence_transformers import SentenceTransformer
import faiss

# BM25
from rank_bm25 import BM25Okapi

from config import settings


# ---------------------------------------------------------------------------
# Supported file extensions
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


# ===========================================================================
# Step 1 — Document Parsing
# ===========================================================================

def _parse_with_unstructured(file_path: Path) -> List[Document]:
    """
    Attempt to parse using `unstructured`. Returns empty list on any failure.
    FIX: Wrapped in try/except per-import so missing extras don't crash everything.
    """
    try:
        from unstructured.partition.auto import partition
        from unstructured.cleaners.core import clean_extra_whitespace
    except ImportError as exc:
        logger.warning(f"unstructured not available ({exc}); falling back to plain-text reader")
        return []

    docs: List[Document] = []
    try:
        elements = partition(
            filename=str(file_path),
            include_page_breaks=True,
            strategy="fast",
        )
    except Exception as exc:
        logger.error(f"unstructured failed on {file_path.name}: {exc}")
        return []

    current_section: str = "Introduction"
    for element in elements:
        category = type(element).__name__
        if category == "Title":
            current_section = str(element).strip()

        text = clean_extra_whitespace(str(element)).strip()
        if len(text) < 30:
            continue

        page_num = None
        if hasattr(element, "metadata") and element.metadata:
            page_num = getattr(element.metadata, "page_number", None)

        metadata = {
            "source": file_path.name,
            "source_path": str(file_path),
            "page": page_num or 1,
            "section": current_section,
            "element_type": category,
            "doc_id": hashlib.md5(file_path.name.encode()).hexdigest()[:8],
        }
        docs.append(Document(page_content=text, metadata=metadata))

    return docs


def _parse_plain_text(file_path: Path) -> List[Document]:
    """
    FIX: Fallback plain-text reader for .txt and .md, and as a last resort
    for PDF/DOCX when unstructured is unavailable.
    """
    docs: List[Document] = []
    try:
        if file_path.suffix.lower() == ".pdf":
            # Try pdfminer
            try:
                from pdfminer.high_level import extract_text as pdf_extract
                text = pdf_extract(str(file_path))
            except Exception:
                logger.error(f"pdfminer failed on {file_path.name}")
                return []
        elif file_path.suffix.lower() in {".docx", ".doc"}:
            try:
                import docx
                doc_obj = docx.Document(str(file_path))
                text = "\n".join(p.text for p in doc_obj.paragraphs)
            except Exception:
                logger.error(f"python-docx failed on {file_path.name}")
                return []
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")

        if not text.strip():
            return []

        metadata = {
            "source": file_path.name,
            "source_path": str(file_path),
            "page": 1,
            "section": "Document",
            "element_type": "PlainText",
            "doc_id": hashlib.md5(file_path.name.encode()).hexdigest()[:8],
        }
        docs.append(Document(page_content=text.strip(), metadata=metadata))
    except Exception as exc:
        logger.error(f"Plain-text fallback failed on {file_path.name}: {exc}")

    return docs


def parse_document(file_path: Path) -> List[Document]:
    """
    Parse a single document. Tries unstructured first, falls back to plain readers.
    """
    logger.info(f"Parsing: {file_path.name}")

    docs = _parse_with_unstructured(file_path)
    if not docs:
        logger.info(f"  → Falling back to plain reader for {file_path.name}")
        docs = _parse_plain_text(file_path)

    logger.info(f"  → {len(docs)} element(s) extracted from {file_path.name}")
    return docs


def load_all_documents(documents_dir: Path) -> List[Document]:
    """
    Recursively find all supported files and parse them.
    Returns a flat list of Documents across all files.
    """
    all_docs: List[Document] = []
    files = [
        f for f in documents_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not files:
        logger.warning(
            f"No supported documents found in '{documents_dir}'. "
            f"Supported types: {SUPPORTED_EXTENSIONS}"
        )
        return all_docs

    logger.info(f"Found {len(files)} document(s) to ingest")
    for file_path in files:
        docs = parse_document(file_path)
        all_docs.extend(docs)

    logger.info(f"Total elements parsed: {len(all_docs)}")
    return all_docs


# ===========================================================================
# Step 2 — Chunking
# ===========================================================================

def chunk_documents(documents: List[Document]) -> List[Document]:
    """
    Split documents into overlapping chunks using RecursiveCharacterTextSplitter.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        add_start_index=True,
    )

    chunks = splitter.split_documents(documents)

    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i
        chunk.metadata["chunk_preview"] = chunk.page_content[:80].replace("\n", " ")

    logger.info(
        f"Chunking complete: {len(documents)} elements → {len(chunks)} chunks "
        f"(size={settings.CHUNK_SIZE}, overlap={settings.CHUNK_OVERLAP})"
    )
    return chunks


# ===========================================================================
# Step 3 — Embeddings
# ===========================================================================

def load_embedding_model() -> SentenceTransformer:
    """Load (or download) the sentence-transformer embedding model."""
    logger.info(f"Loading embedding model: {settings.EMBEDDING_MODEL}")
    model = SentenceTransformer(settings.EMBEDDING_MODEL)
    return model


def embed_chunks(
    chunks: List[Document],
    model: SentenceTransformer,
    batch_size: int = 64,
) -> np.ndarray:
    """
    Generate L2-normalised embeddings for all chunks.
    """
    texts = [chunk.page_content for chunk in chunks]
    logger.info(f"Embedding {len(texts)} chunks with batch_size={batch_size}...")

    embeddings = model.encode(
        texts,
        batch_size=batch_size,
        show_progress_bar=True,
        normalize_embeddings=True,
        convert_to_numpy=True,
    )

    logger.info(f"Embeddings shape: {embeddings.shape}")
    return embeddings.astype(np.float32)


# ===========================================================================
# Step 4 — FAISS Index
# ===========================================================================

def build_faiss_index(embeddings: np.ndarray) -> faiss.Index:
    """Build a FAISS IndexFlatIP (exact inner-product / cosine similarity)."""
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    logger.info(f"FAISS index built — {index.ntotal} vectors, dim={dim}")
    return index


def save_faiss_index(
    index: faiss.Index,
    chunks: List[Document],
    store_path: Path,
) -> None:
    """Persist FAISS index + chunk metadata to disk."""
    store_path.mkdir(parents=True, exist_ok=True)
    faiss.write_index(index, str(store_path / "index.faiss"))

    with open(store_path / "chunks.pkl", "wb") as f:
        pickle.dump(chunks, f)

    logger.info(f"FAISS index persisted to '{store_path}'")


def load_faiss_index(store_path: Path) -> Tuple[faiss.Index, List[Document]]:
    """Load persisted FAISS index and chunk metadata."""
    faiss_path = store_path / "index.faiss"
    chunks_path = store_path / "chunks.pkl"

    # FIX: Raise a clear, actionable error instead of a cryptic file-not-found
    if not faiss_path.exists() or not chunks_path.exists():
        raise FileNotFoundError(
            f"FAISS index not found at '{store_path}'. "
            "Please run ingestion first:\n"
            "  python ingestion.py\n"
            "Or trigger via API: POST /api/ingest"
        )

    index = faiss.read_index(str(faiss_path))
    with open(chunks_path, "rb") as f:
        chunks = pickle.load(f)
    logger.info(f"FAISS index loaded — {index.ntotal} vectors")
    return index, chunks


# ===========================================================================
# Step 5 — BM25 Index (Sparse Retrieval)
# ===========================================================================

def build_bm25_index(chunks: List[Document]) -> BM25Okapi:
    """Build BM25 index from chunk texts (simple whitespace tokenisation)."""
    tokenised = [chunk.page_content.lower().split() for chunk in chunks]
    bm25 = BM25Okapi(tokenised)
    logger.info(f"BM25 index built — {len(tokenised)} documents")
    return bm25


def save_bm25_index(bm25: BM25Okapi, path: Path) -> None:
    with open(path, "wb") as f:
        pickle.dump(bm25, f)
    logger.info(f"BM25 index saved to '{path}'")


def load_bm25_index(path: Path) -> BM25Okapi:
    # FIX: Clear error if BM25 index missing
    if not path.exists():
        raise FileNotFoundError(
            f"BM25 index not found at '{path}'. "
            "Please run ingestion first:\n"
            "  python ingestion.py"
        )
    with open(path, "rb") as f:
        bm25 = pickle.load(f)
    logger.info("BM25 index loaded")
    return bm25


# ===========================================================================
# Orchestrator — run full ingestion pipeline
# ===========================================================================

def run_ingestion(force_rebuild: bool = False) -> None:
    """
    Full ingestion pipeline:
      parse → chunk → embed → FAISS → BM25

    Set force_rebuild=True to re-index even if a saved index exists.
    """
    faiss_ready = (settings.VECTOR_STORE_PATH / "index.faiss").exists()
    bm25_ready = settings.BM25_INDEX_PATH.exists()

    if faiss_ready and bm25_ready and not force_rebuild:
        logger.info(
            "Existing indices found. Skipping ingestion. "
            "Pass --force to rebuild."
        )
        return

    # 1. Parse
    raw_docs = load_all_documents(settings.DOCUMENTS_DIR)
    if not raw_docs:
        raise RuntimeError(
            f"No documents found in '{settings.DOCUMENTS_DIR}'. "
            "Please add HR policy files before ingesting."
        )

    # 2. Chunk
    chunks = chunk_documents(raw_docs)
    if not chunks:
        raise RuntimeError("Chunking produced no output. Check document content.")

    # 3. Embed
    embed_model = load_embedding_model()
    embeddings = embed_chunks(chunks, embed_model)

    # 4. FAISS
    index = build_faiss_index(embeddings)
    save_faiss_index(index, chunks, settings.VECTOR_STORE_PATH)

    # 5. BM25
    bm25 = build_bm25_index(chunks)
    save_bm25_index(bm25, settings.BM25_INDEX_PATH)

    logger.success(
        f"Ingestion complete — {len(chunks)} chunks indexed and ready."
    )


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="HR RAG — Document Ingestion")
    parser.add_argument(
        "--force", action="store_true", help="Force rebuild of indices"
    )
    args = parser.parse_args()

    run_ingestion(force_rebuild=args.force)
